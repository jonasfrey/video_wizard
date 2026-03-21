import docx
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

s_dir = os.path.dirname(os.path.abspath(__file__))
s_path__template = os.path.join(s_dir, 'WEB2_Kurzdokumentation_Vorlage.docx')
s_path__output = os.path.join(s_dir, 'Kurzdokumentation_Video_Wizard.docx')
s_path__gui = os.path.join(s_dir, 'planning', 'gui.jpeg')
s_path__erm = os.path.join(s_dir, 'planning', 'erm.jpeg')

doc = Document(s_path__template)

# Clear all existing content
for p in doc.paragraphs:
    p.clear()
# Remove all paragraphs except the first (can't delete all)
while len(doc.paragraphs) > 1:
    p = doc.paragraphs[-1]
    p._element.getparent().remove(p._element)
# Also remove any tables from template
while len(doc.tables) > 0:
    t = doc.tables[0]
    t._element.getparent().remove(t._element)

# Clear the first paragraph too
doc.paragraphs[0].clear()

# --- Helper functions ---
def f_add_heading(s_text, n_level=1):
    doc.add_heading(s_text, level=n_level)

def f_add_paragraph(s_text, b_bold=False):
    p = doc.add_paragraph()
    run = p.add_run(s_text)
    run.bold = b_bold
    return p

def f_add_table(a_a_s_row, b_header=True):
    table = doc.add_table(rows=len(a_a_s_row), cols=len(a_a_s_row[0]))
    table.style = 'Normal Table'
    for n_idx_row, a_s_cell in enumerate(a_a_s_row):
        for n_idx_col, s_cell in enumerate(a_s_cell):
            cell = table.cell(n_idx_row, n_idx_col)
            cell.text = s_cell
            if b_header and n_idx_row == 0:
                for run in cell.paragraphs[0].runs:
                    run.bold = True
    doc.add_paragraph()  # spacing after table
    return table

# --- Title page ---
p = doc.paragraphs[0]
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Video Wizard')
run.bold = True
run.font.size = Pt(28)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Kurzdokumentation')
run.font.size = Pt(18)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Modul WEB2 \u2013 SPA-Projekt')
run.font.size = Pt(14)

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.add_run('Jonas Immanuel Frey \u2013 jfr159207@stud.gibb.ch')

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.add_run('Abgabedatum: 21.03.2026')

doc.add_page_break()

# --- 1. Projektbeschreibung ---
f_add_heading('1. Projektbeschreibung', 1)

f_add_paragraph(
    'Video Wizard ist eine Single Page Application (SPA) f\u00fcr die Analyse und Komposition von Videomaterial. '
    'Die Applikation erm\u00f6glicht es, aus Videodateien automatisch Audio zu extrahieren, dieses mittels '
    'Spracherkennung (OpenAI Whisper) und Audioklassifikation (Hugging Face Transformers) zu analysieren '
    'und anschliessend aus den erkannten Audio-Events neue Kompositionen zusammenzustellen, die als fertige '
    'Videodateien gerendert werden k\u00f6nnen.'
)

f_add_paragraph(
    'Die Zielgruppe sind Content Creator und Videoproduzenten, die effizient mit bestehendem Videomaterial '
    'arbeiten m\u00f6chten \u2013 etwa um bestimmte Sprachpassagen zu finden, neu zusammenzustellen und als '
    'eigenst\u00e4ndige Clips zu exportieren. Der Mehrwert liegt in der Kombination aus KI-gest\u00fctzter '
    'Audioanalyse und einem intuitiven Kompositions-Editor, der den Workflow vom Rohmaterial zum fertigen '
    'Clip in einer einzigen Anwendung abbildet.'
)

# --- 2. Anforderungen & Zielerreichung ---
f_add_heading('2. Anforderungen & Zielerreichung', 1)

f_add_paragraph('Abgleich mit den Anforderungen und SMART-Zielen aus der Projekteingabe:')

f_add_heading('SMART-Ziele aus der Projekteingabe', 2)

f_add_table([
    ['#', 'SMART-Ziel', 'Status', 'Bemerkung'],
    ['1', 'Video-Upload und Speicherung: Benutzer kann Videodateien \u00fcber die Weboberfl\u00e4che hochladen, die serverseitig gespeichert und in der Datenbank referenziert werden.',
     'Abweichung', 'Statt eines klassischen Uploads wurde ein Dateibrowser implementiert, \u00fcber den Videos direkt aus dem lokalen Dateisystem ausgew\u00e4hlt werden. Da die Applikation lokal l\u00e4uft (localhost), ist ein Upload nicht n\u00f6tig \u2013 der Dateibrowser ist die effizientere L\u00f6sung.'],
    ['2', 'Audio-Extraktion: Audio-Spur eines Videos wird automatisch mittels FFmpeg extrahiert und als separate Datei gespeichert.',
     'Erreicht', 'FFmpeg extrahiert Audio als WAV (16kHz, Mono). Unterst\u00fctzt MP4, MKV, AVI, MOV, WebM.'],
    ['3', 'Speech-to-Text-Analyse: Aus der Audio-Spur werden mithilfe von Whisper alle gesprochenen W\u00f6rter mit Start-/End-Zeitstempeln (ms) erkannt und als Events in der Datenbank gespeichert.',
     'Erreicht', 'Whisper erkennt wortgenaue Zeitstempel. Zus\u00e4tzlich klassifiziert ein Hugging-Face-Transformers-Modell Segmente als Sprache, Musik oder Stille.'],
    ['4', 'Textbasiertes Schneiden: Benutzer kann transkribierten Text sehen und durch Auswahl/Entfernung von W\u00f6rtern das Video schneiden. Resultat wird serverseitig via FFmpeg zusammengeschnitten.',
     'Erreicht', 'Kompositions-Editor mit Wort- und Satz-Ansicht, Volltextsuche, Mehrfachauswahl und FFmpeg-basiertem Rendering.'],
    ['5', 'Weboberfl\u00e4che: Funktionale Weboberfl\u00e4che, \u00fcber die alle Kernfunktionen (Upload, Analyse-Anzeige, Schnitt-Steuerung) bedienbar sind.',
     'Erreicht', 'Vue-3-SPA mit 5 Ansichten: Data, Filebrowser, Composition, Export, CLI Monitor.'],
])

f_add_heading('Geplante Features aus der Projekteingabe', 2)

f_add_table([
    ['Feature', 'Status', 'Bemerkung'],
    ['Video-Upload \u00fcber die Weboberfl\u00e4che', 'Abweichung', 'Ersetzt durch Dateibrowser (siehe SMART-Ziel 1)'],
    ['Automatische Extraktion der Audio-Spur (FFmpeg)', 'Erreicht', '\u2013'],
    ['Speech-to-Text-Analyse mit wortgenauen Zeitstempeln (Whisper)', 'Erreicht', 'Zus\u00e4tzlich Audioklassifikation (Musik/Stille) implementiert'],
    ['Darstellung des transkribierten Textes mit Zeitstempel-Zuordnung', 'Erreicht', 'Wort- und Satz-Ansicht mit farbiger Typ-Kennzeichnung'],
    ['Textbasiertes Schneiden: Auswahl/Entfernung von W\u00f6rtern', 'Erreicht', 'Inklusive Vorschau im Browser vor dem Rendern'],
    ['Export des geschnittenen Videos', 'Erreicht', 'MP4-Download mit Dateigr\u00f6sse und Render-Status'],
    ['Datenmodell mit Entit\u00e4ten (Video, Audio, Event, etc.)', 'Erreicht', '12 Modelle implementiert; Object, Action, Sense (f\u00fcr zuk\u00fcnftige Erweiterung) noch nicht umgesetzt'],
    ['Echtzeit-Kommunikation via WebSockets', 'Erreicht', 'Bidirektionale Synchronisation mit Auto-Reconnect'],
])

f_add_heading('Zus\u00e4tzlich implementierte Features', 2)

a_s_extra = [
    'CLI-Monitor: Echtzeit-\u00dcberwachung aller laufenden Prozesse (Extraktion, Analyse, Rendering)',
    'Satz-Gruppierung: Automatische Gruppierung von Wort-Events zu S\u00e4tzen im Kompositions-Editor',
    'L\u00fccken-Erkennung: Automatische Erkennung nicht-klassifizierter Abschnitte zwischen Events',
    'Per-Event-Vorschau: Einzelne Audio-Events k\u00f6nnen direkt im Editor abgespielt werden',
    'CRUD-Interface: Generisches Datenmanagement f\u00fcr alle Modelle',
]
for s in a_s_extra:
    f_add_paragraph('\u2022  ' + s)

doc.add_paragraph()
f_add_paragraph(
    'Von den 5 SMART-Zielen wurden 4 vollst\u00e4ndig erreicht. SMART-Ziel 1 (Video-Upload) wurde bewusst '
    'durch einen Dateibrowser ersetzt, da die Applikation lokal l\u00e4uft und ein Upload-Mechanismus unn\u00f6tig '
    'w\u00e4re. Alle 8 geplanten Features wurden umgesetzt, wobei der Video-Upload durch den Dateibrowser '
    'abgel\u00f6st wurde und die Entit\u00e4ten Object, Action und Sense als geplante Erweiterung noch ausstehen. '
    'Dar\u00fcber hinaus wurden 5 zus\u00e4tzliche Features implementiert, die \u00fcber den urspr\u00fcnglichen '
    'Projektumfang hinausgehen.'
)

# --- 3. Technische Umsetzung ---
f_add_heading('3. Technische Umsetzung', 1)

f_add_heading('3.1 Voraussetzungen', 2)

f_add_table([
    ['Abh\u00e4ngigkeit', 'Version', 'Zweck'],
    ['Deno', '2.x', 'Server-Runtime'],
    ['Python', '3.9+', 'Audioanalyse- und Rendering-Scripts'],
    ['FFmpeg / FFprobe', 'beliebig', 'Audio-/Videoextraktion und Rendering'],
    ['pip3', 'beliebig', 'Python-Paketmanager'],
])

f_add_paragraph('Python-Pakete (werden im Virtual Environment installiert):', b_bold=True)
a_s_pypackage = [
    'openai-whisper \u2014 Speech-to-Text-Modell',
    'torch (PyTorch) \u2014 ML-Framework',
    'transformers \u2014 Hugging Face Audioklassifikation',
    'numpy \u2014 Numerische Berechnungen',
    'pyttsx3 \u2014 Text-to-Speech-Synthese',
    'python-dotenv \u2014 Umgebungsvariablen laden',
]
for s in a_s_pypackage:
    f_add_paragraph('\u2022  ' + s)

doc.add_paragraph()
f_add_paragraph(
    'GPU nicht erforderlich. PyTorch, Whisper und Transformers laufen standardm\u00e4ssig auf der CPU. '
    'Eine NVIDIA-GPU mit CUDA-Unterst\u00fctzung beschleunigt die Audioanalyse, ist aber keine '
    'Voraussetzung. Im Code wird kein explizites GPU-Device gesetzt \u2014 die Inferenz findet auf '
    'dem von PyTorch gew\u00e4hlten Standard-Device statt (CPU).'
)

f_add_heading('3.2 Architektur', 2)

f_add_paragraph(
    'Die Applikation basiert auf einer Client-Server-Architektur mit folgenden Technologien:'
)

a_s_tech = [
    'Frontend: Vue 3 (ESM-Module, kein Build-Step), Vue Router 4 (Hash-basiertes Routing)',
    'Backend: Deno (TypeScript-Runtime) mit WebSocket-Server',
    'Datenbank: SQLite oder JSON (konfigurierbar \u00fcber Umgebungsvariable S_DB_TYPE)',
    'CLI-Tools: FFmpeg/FFprobe (Audioextraktion), Python 3 (Audioanalyse und Rendering)',
    'Python-Bibliotheken: OpenAI Whisper, Hugging Face Transformers, PyTorch, pyttsx3',
]
for s in a_s_tech:
    f_add_paragraph('\u2022  ' + s)

doc.add_paragraph()
f_add_paragraph(
    'Die gesamte Kommunikation zwischen Client und Server erfolgt \u00fcber WebSocket. HTTP wird ausschliesslich '
    'f\u00fcr das Ausliefern statischer Dateien und Mediendateien verwendet. Der Server verwaltet einen reaktiven '
    'o_state, der bei jeder \u00c4nderung automatisch an alle verbundenen Clients propagiert wird.'
)

f_add_paragraph('Projektstruktur:', b_bold=True)

# Use a code-like paragraph for structure
s_struct = (
    'video_wizard/\n'
    '\u251c\u2500\u2500 localhost/                  # Frontend (Vue 3 SPA)\n'
    '\u2502   \u251c\u2500\u2500 index.html              # HTML-Shell\n'
    '\u2502   \u251c\u2500\u2500 index.js                # Vue-App, Routing, WebSocket\n'
    '\u2502   \u251c\u2500\u2500 index.css               # Globales Styling (Dark Theme)\n'
    '\u2502   \u251c\u2500\u2500 constructors.js         # Datenmodelle & WebSocket-Nachrichten\n'
    '\u2502   \u251c\u2500\u2500 o_component__*.js       # 5 Hauptkomponenten\n'
    '\u2502   \u2514\u2500\u2500 lib/                    # Vue 3 & Vue Router ESM-Bundles\n'
    '\u251c\u2500\u2500 serverside/                 # Backend (Deno)\n'
    '\u2502   \u251c\u2500\u2500 cli_functions.js        # CLI-Subprozesse (FFmpeg, Python)\n'
    '\u2502   \u251c\u2500\u2500 database_functions.js   # SQLite-Operationen\n'
    '\u2502   \u251c\u2500\u2500 f_analyze_audio.py      # Spracherkennung & Klassifikation\n'
    '\u2502   \u251c\u2500\u2500 f_render_composition.py # Video-Rendering\n'
    '\u2502   \u2514\u2500\u2500 testing/                # Unit-Tests\n'
    '\u251c\u2500\u2500 deno.json                   # Deno-Konfiguration & Tasks\n'
    '\u251c\u2500\u2500 .env                        # Umgebungsvariablen\n'
    '\u2514\u2500\u2500 readme.md                   # Setup-Anleitung'
)
p = doc.add_paragraph()
run = p.add_run(s_struct)
run.font.size = Pt(8)
run.font.name = 'Courier New'

f_add_heading('3.3 Komponenten', 2)

f_add_paragraph(
    'Die Applikation besteht aus f\u00fcnf Hauptkomponenten, die jeweils eine eigene Verantwortlichkeit haben:'
)

f_add_table([
    ['Komponente', 'Datei', 'Beschreibung'],
    ['Data', 'o_component__data.js', 'CRUD-Interface f\u00fcr alle Datenmodelle (Erstellen, Bearbeiten, L\u00f6schen)'],
    ['Filebrowser', 'o_component__filebrowser.js', 'Dateisystem-Navigation, Videoauswahl, Analyse-Trigger'],
    ['Composition', 'o_component__composition.js', 'Audio-Event-Auswahl, Volltextsuche, Kompositions-Vorschau'],
    ['Export', 'o_component__export.js', 'Rendering-Ausl\u00f6sung, Download, Vorschau gerenderter Videos'],
    ['CLI Monitor', 'o_component__cli_monitor.js', 'Echtzeit-\u00dcberwachung aller laufenden CLI-Tasks'],
])

f_add_paragraph(
    'Zus\u00e4tzlich gibt es \u00fcbergreifende Module: constructors.js definiert alle 12 Datenmodelle und '
    'WebSocket-Nachrichtentypen, functions.js enth\u00e4lt gemeinsame Hilfsfunktionen (Client & Server), '
    'und bgshader.js steuert den animierten Canvas-Hintergrund. Die Kommunikation zwischen Komponenten '
    'erfolgt \u00fcber den gemeinsamen reaktiven o_state und WebSocket-Nachrichten.'
)

f_add_paragraph(
    'Der gesamte Code folgt der Abstract Prefix Notation (APN) \u2014 einer Namenskonvention, bei der '
    'der Datentyp als Pr\u00e4fix im Variablennamen kodiert wird (z.B. s_name f\u00fcr Strings, n_age '
    'f\u00fcr Zahlen, a_o_video f\u00fcr Arrays von Objekten, f_render f\u00fcr Funktionen). Dies '
    'erh\u00f6ht die Lesbarkeit und macht den Code selbstdokumentierend. '
    'Das Whitepaper zur APN-Methodik ist unter techrxiv.org ver\u00f6ffentlicht.'
)

f_add_heading('3.4 Routing', 2)

f_add_paragraph(
    'Das Routing ist mit Vue Router 4 und Hash-basierter Navigation (createWebHashHistory) implementiert:'
)

f_add_table([
    ['Route', 'Komponente', 'Beschreibung'],
    ['/', '\u2013', 'Redirect auf /data'],
    ['/data', 'Data', 'Datenmanagement aller Modelle'],
    ['/filebrowser', 'Filebrowser', 'Dateisystem-Navigation & Videoanalyse'],
    ['/composition', 'Composition', 'Audio-Event-Komposition'],
    ['/export', 'Export', 'Video-Rendering & Download'],
    ['/cli', 'CLI Monitor', 'Prozess\u00fcberwachung'],
])

f_add_paragraph(
    'Die Navigation erfolgt \u00fcber eine persistente Navigationsleiste. Deep Links funktionieren dank '
    'Hash-Routing. Die zuletzt besuchte Seite wird in der Datenbank gespeichert und beim n\u00e4chsten '
    'Start wiederhergestellt.'
)

f_add_heading('3.5 State-Management', 2)

f_add_paragraph(
    'Anstelle eines klassischen Store-Frameworks (Pinia, Vuex) wird ein eigenes reaktives '
    'State-Management-System verwendet. Diese bewusste Entscheidung basiert auf folgenden Gr\u00fcnden:'
)

a_s_state = [
    'Echtzeit-Synchronisation: Der State wird \u00fcber WebSocket bidirektional zwischen Server und allen '
    'Clients synchronisiert. Ein klassischer clientseitiger Store w\u00fcrde diese Architektur verkomplizieren.',
    'Einheitliches Datenmodell: o_state ist ein Vue-3-reactive-Objekt, das alle Datenmodelle als Arrays '
    'enth\u00e4lt (z.B. o_state.a_o_video, o_state.a_o_composition). CRUD-Operationen werden \u00fcber die '
    'zentrale Funktion f_v_sync() ausgef\u00fchrt, die gleichzeitig die Datenbank, den Server-State und alle '
    'verbundenen Clients aktualisiert.',
    'Automatische Denormalisierung: Fremdschl\u00fcssel-Beziehungen werden automatisch aufgel\u00f6st '
    '(z.B. o_video.o_fsnode wird aus n_o_fsnode_n_id denormalisiert).',
]
for s in a_s_state:
    f_add_paragraph('\u2022  ' + s)

doc.add_paragraph()
f_add_paragraph('State-Bereiche:', b_bold=True)
f_add_paragraph('\u2022  Anwendungsdaten: Alle Modelle (Videos, Audios, Audio-Events, Kompositionen, etc.)')
f_add_paragraph('\u2022  UI-State: Ladefortschritt, Toast-Nachrichten, CLI-Task-Status, aktive Route')

# --- 4. Visualisierungen ---
f_add_heading('4. Visualisierungen', 1)

f_add_heading('4.1 GUI-Entwurf (Wireframe)', 2)

f_add_paragraph(
    'Der folgende handgezeichnete Wireframe zeigt den fr\u00fchen GUI-Entwurf der Applikation. '
    'Er zeigt die Grundstruktur mit Navigationsleiste, Videoanzeige und den verschiedenen '
    'Bedienelementen f\u00fcr die Videoanalyse und den Kompositions-Editor.'
)

doc.add_picture(s_path__gui, width=Inches(5.5))
last_paragraph = doc.paragraphs[-1]
last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Abbildung 1: Fr\u00fcher GUI-Entwurf (Wireframe) der Video Wizard Applikation')
run.italic = True
run.font.size = Pt(9)

f_add_heading('4.2 Entity-Relationship-Modell (ERM)', 2)

f_add_paragraph(
    'Das Entity-Relationship-Modell zeigt die Datenstruktur der Applikation. Die zentralen Entit\u00e4ten '
    'sind fsnode (Dateisystem-Knoten), Video, Audio, AudioEvent und Composition. '
    'Ein fsnode verweist auf eine Videodatei, aus der ein Audio extrahiert wird. Das Audio wird in '
    'einzelne AudioEvents aufgeteilt (Sprache, Musik, Stille). Der Benutzer w\u00e4hlt AudioEvents aus '
    'und fasst sie zu einer Composition zusammen, die als neues Video gerendert werden kann.'
)

doc.add_picture(s_path__erm, width=Inches(5.5))
last_paragraph = doc.paragraphs[-1]
last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Abbildung 2: Entity-Relationship-Modell der Video Wizard Datenstruktur')
run.italic = True
run.font.size = Pt(9)

f_add_heading('4.3 Screenshots der fertigen Applikation', 2)

s_path__gui_dir = os.path.join(s_dir, 'gui', 'v1')

a_a_s_screenshot = [
    ['filebrowser_page.png', 'Abbildung 3: Filebrowser \u2014 Dateisystem-Navigation mit analysierten Videos in der Seitenleiste'],
    ['composition_page.png', 'Abbildung 4: Kompositions-Editor \u2014 Videovorschau, Transkript-Suche und Audio-Event-Auswahl'],
    ['export_page.png', 'Abbildung 5: Export \u2014 Gerenderte Kompositionen mit Vorschau und Download'],
    ['cli_page.png', 'Abbildung 6: CLI Monitor \u2014 Echtzeit-\u00dcberwachung laufender Prozesse'],
]

for a_s in a_a_s_screenshot:
    s_path_img = os.path.join(s_path__gui_dir, a_s[0])
    doc.add_picture(s_path_img, width=Inches(5.5))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(a_s[1])
    run.italic = True
    run.font.size = Pt(9)

f_add_heading('4.4 Architektur\u00fcbersicht', 2)

f_add_paragraph(
    'Die folgende \u00dcbersicht zeigt den Datenfluss der Video-Analysepipeline:'
)

f_add_table([
    ['Schritt', 'Prozess', 'Technologie', 'Ergebnis'],
    ['1', 'Videodatei ausw\u00e4hlen', 'Dateibrowser (Vue 3)', 'o_fsnode + o_video in DB'],
    ['2', 'Audio extrahieren', 'FFmpeg (serverseitig)', 'WAV-Datei + o_audio in DB'],
    ['3', 'Audio analysieren', 'Whisper + Transformers (Python)', 'o_audio_event Eintr\u00e4ge in DB'],
    ['4', 'Komposition erstellen', 'Kompositions-Editor (Vue 3)', 'o_composition + Zuordnungstabelle'],
    ['5', 'Video rendern', 'FFmpeg (serverseitig)', 'Fertiges MP4-Video'],
])

# --- 5. Fazit & Lessons Learned ---
f_add_heading('5. Fazit & Lessons Learned', 1)

f_add_heading('Herausforderungen', 2)

# Challenge 1
p = doc.add_paragraph()
run = p.add_run('1. Spracherkennung und Sprachklassifikation: ')
run.bold = True
p.add_run(
    'Bei einem Testvideo wurde die Sprache f\u00e4lschlicherweise als Chinesisch erkannt, was die gesamte '
    'Transkription unbrauchbar machte. Die Ursache liegt im Whisper-Modell, das bei kurzen oder undeutlichen '
    'Audiopassagen die Sprache falsch identifizieren kann. Als L\u00f6sung wurde ein --s-language Parameter '
    'implementiert, der die Sprache manuell vorgeben l\u00e4sst, anstatt sich vollst\u00e4ndig auf die '
    'automatische Erkennung zu verlassen.'
)

# Challenge 2
p = doc.add_paragraph()
run = p.add_run('2. Richtiges Prompting mit KI-Werkzeugen: ')
run.bold = True
p.add_run(
    'Die Entwicklung erfolgte unter starkem Einsatz von Claude Code als KI-Assistenten. Eine zentrale '
    'Herausforderung war das pr\u00e4zise Formulieren von Anweisungen (Prompts). Oft programmierte die KI '
    'nicht exakt nach Vorstellung oder liess gew\u00fcnschte Funktionalit\u00e4ten weg. Es war schwierig '
    'nachzuvollziehen, was im Hintergrund mit den Prompts geschieht und wie diese verarbeitet werden. '
    'Richtiges Prompting ist eine F\u00e4higkeit, die erlernt und ge\u00fcbt werden muss.'
)

# Challenge 3
p = doc.add_paragraph()
run = p.add_run('3. Echtzeit-Synchronisation \u00fcber WebSocket: ')
run.bold = True
p.add_run(
    'Die bidirektionale Echtzeit-Synchronisation des Anwendungsstatus \u00fcber WebSocket stellte sich als '
    'komplex heraus. Die Herausforderung bestand darin, CRUD-Operationen konsistent auf Server, Datenbank '
    'und allen verbundenen Clients gleichzeitig auszuf\u00fchren, ohne Race Conditions zu erzeugen. Die '
    'L\u00f6sung war eine zentrale f_v_sync()-Funktion, die als einziger Eingangspunkt f\u00fcr alle '
    'Daten\u00e4nderungen dient.'
)

f_add_heading('Pers\u00f6nliche Erkenntnisse', 2)

f_add_paragraph(
    'Das Arbeiten mit KI ist durchaus praktisch und beschleunigt die Entwicklung erheblich. Dennoch '
    'ersetzt es nicht das Verst\u00e4ndnis der zugrunde liegenden Technologien. Es ist immer wieder '
    'spannend zu sehen, wie m\u00e4chtig kontextbewusste KI-Werkzeuge sind, aber man muss lernen, sie '
    'effektiv einzusetzen. Die Kombination aus eigenem Fachwissen und KI-Unterst\u00fctzung ergibt '
    'das beste Resultat.'
)

f_add_heading('Aussicht & Verbesserungsvorschl\u00e4ge', 2)

f_add_paragraph(
    'Die Applikation kann zu einem vollst\u00e4ndigen Video-Editor erweitert werden. '
    'Konkrete n\u00e4chste Schritte w\u00e4ren:'
)

a_s_improve = [
    'FFmpeg im Browser: F\u00fcr eine rein clientseitige L\u00f6sung k\u00f6nnte FFmpeg mit WebAssembly '
    '(ffmpeg.wasm) direkt im Browser ausgef\u00fchrt werden, was den Server-Prozess eliminieren w\u00fcrde.',
    'Optische Videoanalyse: Neben der Audioanalyse k\u00f6nnte eine Bildanalyse (z.B. Szenerkennung, '
    'Objekterkennung) zus\u00e4tzliche Kontextinformationen f\u00fcr den Editor liefern.',
    'Manuelle Sprachauswahl: Um das Problem der falschen Spracherkennung zu umgehen, k\u00f6nnte eine '
    'manuelle Spracheinstellung als Fallback angeboten werden.',
    'Detaillierteres Fehlerhandling: Bei fehlgeschlagenen Analysen k\u00f6nnte dem Benutzer genauer '
    'erkl\u00e4rt werden, warum der Prozess gescheitert ist und welche Schritte zur Behebung m\u00f6glich sind.',
]
for s in a_s_improve:
    f_add_paragraph('\u2022  ' + s)

# --- 6. Quellen ---
f_add_heading('6. Quellen', 1)

a_a_s_source = [
    ['Video Wizard Repository', 'https://github.com/jonasfrey/video_wizard'],
    ['Deno', 'https://github.com/denoland/deno'],
    ['FFmpeg', 'https://github.com/FFmpeg/FFmpeg'],
    ['OpenAI Whisper', 'https://github.com/openai/whisper'],
    ['PyTorch', 'https://github.com/pytorch/pytorch'],
    ['Hugging Face Transformers', 'https://github.com/huggingface/transformers'],
    ['Vue 3', 'https://github.com/vuejs/core'],
    ['Vue Router 4', 'https://github.com/vuejs/router'],
    ['pyttsx3', 'https://github.com/nateshmbhat/pyttsx3'],
    ['python-dotenv', 'https://github.com/theskumar/python-dotenv'],
    ['NumPy', 'https://github.com/numpy/numpy'],
    ['Abstract Prefix Notation (APN)', 'https://www.techrxiv.org/users/1031649/articles/1391488-abstract-prefix-notation-apn-a-type-encoding-naming-methodology-for-programming'],
]

f_add_table([['Tool / Bibliothek', 'URL']] + a_a_s_source)

# --- Page numbers in footer ---
for section in doc.sections:
    footer = section.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    fld_char_begin = OxmlElement('w:fldChar')
    fld_char_begin.set(qn('w:fldCharType'), 'begin')
    run._element.append(fld_char_begin)

    instr_text = OxmlElement('w:instrText')
    instr_text.set(qn('xml:space'), 'preserve')
    instr_text.text = ' PAGE '
    run._element.append(instr_text)

    fld_char_end = OxmlElement('w:fldChar')
    fld_char_end.set(qn('w:fldCharType'), 'end')
    run._element.append(fld_char_end)

    run.font.size = Pt(9)

# Save
doc.save(s_path__output)
print(f'Dokument gespeichert: {s_path__output}')
